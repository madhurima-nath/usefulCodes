# get necessary libraries
import extract_contents
import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from itertools import zip_longest, chain


def pair_equal_date_url_cols(d):
    """
    returns a list of tuple for
    updated dates, urls, information
    in previous two columns when
    each date corresponds to each
    updated link
    """
    col_date, col_url, col_prev1, col_prev2 = [], [], [], []
    for x in d.get("prev_col1"):
        col_prev1.append(f"Information in Title column:\n{x}")
    for x in d.get("prev_col2"):
        col_prev2.append(f"Information in Sub-title column:\n{x}")
    for x in d.get("date"):
        if type(x) == datetime.datetime:
            vdate = x.strftime("%Y-%b-%d")
        else:
            vdate = x
        if x == "":
            col_date.append(f"DATE(S):\n{vdate}")
    if "url" in d.keys():
        if len(d.get("url")) > 0:
            for x in d.get("url"):
                col_url.append(f"LINK(S):\n{x}")
    if col_date and col_url:
        result = list(zip(col_date, col_url, col_prev2, col_prev1))
    elif len(col_date) > len(col_url):
        result = list(zip_longest(col_date, col_url, col_prev2, col_prev1,
                                    fillvalue="Links not updated (CHECK)."))
    else:
        result = list(zip(["Missing date in Excel (CHECK)"],
                            ["Links not updated (CHECK)"],
                            col_prev2, col_prev1))
    return result


def pair_random_date_url_cols(d):
    """
    returns a list of tuples for
    updated dates, urls, information
    in previous two columns when
    each date does not correspond
    to updated links
    """
    col_prev1, col_prev2 = [], []
    for x in d.get("prev_col1"):
        col_prev1.append(f"Information in Title column:\n{x}")
    for x in d.get("prev_col2"):
        col_prev2.append(f"Information in Sub-title column:\n{x}")
    for x in d.get("date"):
        temp_url = []
        if type(x) == str:
            num = len(x.split("\n"))
        url_list = d.get("url")[:num]
        for link in url_list:
            temp_url.append(f"\nLINK(S):\n\n{link}\n\n")
        temp_date = []
        if type(x) == datetime.datetime:
            vdate = x.strftime("%Y-%b-%d")
            temp_date.append(f"DATE(S):\n{vdate}")
        else:
            for date in x.split("\n"):
                vdate = date.strip()
                temp_date.append(f"DATE(S):\n{vdate}")
        
        temp_list = list(zip_longest(temp_date, temp_url))
        result = list(chain(*zip(temp_list, col_prev2, col_prev1)))


def word_output(df):
    """
    this function writes the output
    to a word document.
    since every country has a different
    input structure, multiple if conditions
    are used to capture these, along with
    a generic section
    """

    # create a new document
    document = Document()

    # get output
    result = {}
    for x in df["country"].unique():
        temp = extract_contents.get_updates(df, x)
        if temp and x not in result:
            result = temp
        try:
            if "pdf" in result[x]["url"][0]:
                url = result[x]["url"][0]
                search_text = result[x]["search_text"]
                result = extract_contents.get_pdf_contents(url, search_text)
        except:
            print(f"No pdfs updated for {x}. Only updated urls and/or dates.")
            result = temp

        print(result)

        for key, value in result.items():
            # add country as heading
            country = value.get("country")
            heading = document.add_heading(country, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            heading.runs[0].font.color.rgb = RGBColor(255, 0, 0) # red color
            heading.runs[0].font.underline = True # underline

            # start document body
            document.add_paragraph("Updates/Changes:".upper())

            if len(value.get("date")) < len(value.get("url")):
                pair_cols = pair_random_date_url_cols(value)
                document.add_paragraph("information on dates, links and previous columns:".upper())
                for x in pair_cols:
                    document.add_paragraph(x)
            else:
                pair_cols = pair_equal_date_url_cols(value)
                if len(value.get("date")) == len(pair_cols) == len(value.get("url")):
                    document.add_paragraph("information on dates, links and previous columns:".upper())
                    for x in pair_cols:
                        document.add_paragraph("\n\n".join(x))
                elif len(value.get("date")) == len(pair_cols) > len(value.get("url")):
                    document.add_paragraph("there is a misalignment between updated dates and/or links.".upper())
                    document.add_paragraph("or".upper())
                    document.add_paragraph("some updates might be for dates only, not links.".upper())
                    for x in pair_cols:
                        document.add_paragraph("\n\n".join(x))
                elif len(value.get("date")) == len(pair_cols):
                    document.add_paragraph("all updates are for dates only, not links.".upper())
                    for x in pair_cols:
                        document.add_paragraph("\n\n".join(x))

            document.add_paragraph("Comments:".upper())
            document.add_paragraph(f"{value.get('comments')}")

    return document